from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# ── Helpers ─────────────────────────────────────────────────────────
def rgb(hex_str):
    h = hex_str.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)

def styled_cell(cell, text, bg='FFFFFF', bold=False, font_size=9, color='000000', align='left', italic=False):
    set_cell_bg(cell, bg)
    set_cell_borders(cell)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.clear()
    al = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = al.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.color.rgb = rgb(color)
    run.font.name = 'Arial'

def set_col_widths(table, widths_cm):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(sum(widths_cm) * 567)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(widths_cm[j] * 567)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb('1F3864')
    run.font.name = 'Arial'
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb('2E5C8A')
    run.font.name = 'Arial'

def body(text, bold=False, size=9, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Arial'

def spacer(pts=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)

# ── Title block ──────────────────────────────────────────────────────
t = doc.add_table(rows=2, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.CENTER

c = t.cell(0, 0)
set_cell_bg(c, '1F3864')
set_cell_borders(c)
set_cell_margins(c, top=140, bottom=100, left=200, right=200)
c.paragraphs[0].clear()
c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
r = c.paragraphs[0].add_run('Round 1 Adapted Report')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = rgb('FFFFFF'); r.font.name = 'Arial'
p2 = c.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Single User Prompt  |  Category Difficulty Bands')
r2.font.size = Pt(11); r2.font.color.rgb = rgb('BDD7EE'); r2.font.name = 'Arial'

c2 = t.cell(1, 0)
set_cell_bg(c2, 'EAF1FB')
set_cell_borders(c2)
set_cell_margins(c2, top=60, bottom=60, left=200, right=200)
c2.paragraphs[0].clear()
c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = c2.paragraphs[0].add_run('Models: claude-opus-4-8, claude-sonnet-4-6, claude-haiku-4-5  |  11 categories  |  7 tiers  |  Finance domain  |  Adapted from capability_testing.xlsx')
r3.font.size = Pt(8); r3.font.color.rgb = rgb('444444'); r3.font.name = 'Arial'

set_col_widths(t, [18.0])
spacer(14)

# ── Section 1: Single User Prompt ───────────────────────────────────
heading1('1. Single User Prompt')

body('The original Round 4 test used 8 varying user prompts (UP1-UP8) across different ticker pairs and data sources. This adaptation replaces all 8 with one fixed prompt so that system prompt complexity (tier) is the only variable being tested.')
spacer(6)

# Prompt box
tp = doc.add_table(rows=1, cols=1)
tp.alignment = WD_TABLE_ALIGNMENT.CENTER
cp = tp.cell(0,0)
set_cell_bg(cp, 'EBF5FB')
set_cell_borders(cp)
set_cell_margins(cp, top=120, bottom=120, left=200, right=200)
cp.paragraphs[0].clear()
lbl = cp.paragraphs[0].add_run('Fixed User Prompt')
lbl.bold = True; lbl.font.size = Pt(9); lbl.font.color.rgb = rgb('1F3864'); lbl.font.name = 'Arial'
p_prompt = cp.add_paragraph()
r_prompt = p_prompt.add_run('"Using only the attached data, summarise the most recent quarterly earnings performance, insider share activity, and daily price movement for Microsoft Corp (MSFT) and JPMorgan Chase & Co (JPM)."')
r_prompt.italic = True; r_prompt.font.size = Pt(9.5); r_prompt.font.name = 'Arial'
p_src = cp.add_paragraph()
r_src = p_src.add_run('Data sources required: earnings_and_insider.csv  |  prices.csv')
r_src.font.size = Pt(8); r_src.font.color.rgb = rgb('666666'); r_src.font.name = 'Arial'
set_col_widths(tp, [18.0])

spacer(8)
body('Ticker selection rationale: MSFT and JPM are both confirmed present in the DJIA dataset (referenced in UP1 of the original Round 4 design). They represent contrasting sectors — technology vs. financials — which increases the probability that one ticker beats EPS estimates while the other shows distinct price volatility. This ensures the Situation and Chained conditional rules are genuinely triggered rather than vacuously satisfied across all tiers.', size=8.5)
spacer(6)

heading2('Rule coverage')
rc = doc.add_table(rows=5, cols=3)
rc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(rc, [3.8, 3.2, 11.0])
rows_data = [
    ('Category', 'Rule Type', 'Why This Prompt Fires It', '1F3864', True, 'FFFFFF'),
    ('Sco — Scope', 'Grounding constraint', '"Using only the attached data" directly triggers the scope rule. The model cannot use prior knowledge, infer values, or draw conclusions not in the dataset.', 'F2F7FC', True, '000000'),
    ('Neg — Negation', 'Forbidden words', 'Describing earnings performance and price movement creates natural opportunities to use forbidden words (strong, robust, significant, notable, approximately).', 'FFFFFF', True, '000000'),
    ('Sit — Situation', 'Conditional (if/then)', '"Earnings performance" asks about EPS vs. estimate — triggers the Beat: prefix rule and the T7 significant miss rule whenever conditions are met in the data.', 'F2F7FC', True, '000000'),
    ('Chn — Chained', 'Two-condition logic', '"Daily price movement" + "insider share activity" together provide diff_pct and netShChg — both required for the chained condition (diff_pct > 5% AND netShChg negative, same ticker).', 'FFFFFF', True, '000000'),
]
for i, (c1, c2, c3, bg, hdr_bold, hdr_color) in enumerate(rows_data):
    is_header = (i == 0)
    bg_use = bg
    styled_cell(rc.cell(i,0), c1, bg=bg_use, bold=is_header or True, color='FFFFFF' if is_header else '000000')
    styled_cell(rc.cell(i,1), c2, bg=bg_use, bold=is_header, color='FFFFFF' if is_header else '000000')
    styled_cell(rc.cell(i,2), c3, bg=bg_use, bold=is_header, color='FFFFFF' if is_header else '000000')

spacer(14)

# ── Section 2: Difficulty Bands ──────────────────────────────────────
heading1('2. Category Difficulty Bands')

body('Each category is assigned a difficulty band — Easy, Medium, or Hard — based on its weight in the Round 1 results. Weights are derived empirically from observed failure rates across all three models and all seven tiers. They are not assigned by judgment.')
spacer(6)

# Formula box
tf = doc.add_table(rows=1, cols=1)
tf.alignment = WD_TABLE_ALIGNMENT.CENTER
cf = tf.cell(0,0)
set_cell_bg(cf, 'F8F9FA')
set_cell_borders(cf)
set_cell_margins(cf, top=100, bottom=100, left=160, right=160)
cf.paragraphs[0].clear()
r_lbl = cf.paragraphs[0].add_run('Weight formula')
r_lbl.bold = True; r_lbl.font.size = Pt(9); r_lbl.font.name = 'Arial'
p_f = cf.add_paragraph()
r_f = p_f.add_run('Weight  =  total failures across all models  /  (models x tiers)  =  (Opus + Sonnet + Haiku fails)  /  21')
r_f.font.size = Pt(8.5); r_f.font.name = 'Courier New'
p_note = cf.add_paragraph()
r_note = p_note.add_run('Range: 0 = no model ever failed.  1 = every model failed every tier.')
r_note.font.size = Pt(8); r_note.font.color.rgb = rgb('555555'); r_note.font.name = 'Arial'
set_col_widths(tf, [18.0])
spacer(8)

heading2('Band thresholds')
tb = doc.add_table(rows=5, cols=3)
tb.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tb, [2.8, 3.2, 12.0])
bands = [
    ('Band', 'Weight Threshold', 'Rationale', '1F3864', 'FFFFFF', True),
    ('Hard', 'weight >= 0.50', 'Category fails the majority of runs across models. Scope (0.71) is the only Hard category — all three models failed from T2 onward, making it the primary driver of the weighted score penalty.', 'FADADD', '000000', False),
    ('Medium', '0.20 <= weight < 0.50', 'Category fails roughly one-third to half of all runs. Conditional and chained logic (Situation 0.38, Chained 0.38) fall here — models handle them at low tier counts but degrade under rule accumulation.', 'FFF3CD', '000000', False),
    ('Easy', '0 < weight < 0.20', 'Category rarely fails. Word-level constraints (Negation 0.19, Persona 0.10, Format/Example/Precision 0.05) are tractable even at T7. Included in weighted scoring but with low penalty.', 'D4EDDA', '000000', False),
    ('Excluded', 'weight = 0', 'No model ever failed. Content, Style, and Priority contributed zero failures and carry no discriminative value for cross-category testing.', 'F0F0F0', '888888', False),
]
for i, (c1, c2, c3, bg, txt_col, bold) in enumerate(bands):
    is_hdr = (i == 0)
    styled_cell(tb.cell(i,0), c1, bg=bg, bold=True, color='FFFFFF' if is_hdr else txt_col)
    styled_cell(tb.cell(i,1), c2, bg=bg, bold=is_hdr, color='FFFFFF' if is_hdr else txt_col)
    styled_cell(tb.cell(i,2), c3, bg=bg, bold=is_hdr, color='FFFFFF' if is_hdr else txt_col)

spacer(10)

heading2('All 11 categories with bands')
tc = doc.add_table(rows=8, cols=6)
tc.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tc, [3.8, 2.0, 1.8, 2.0, 2.2, 6.2])
cat_rows = [
    ('Category', 'Fails / 21', 'Weight', 'Band', 'Platform Score', 'Notes', '1F3864', 'FFFFFF', True),
    ('Sco — Scope', '15 / 21', '0.71', 'Hard', '29%', 'All models failed from T2. Highest-impact category.', 'FADADD', '8B0000', False),
    ('Sit — Situation', '8 / 21', '0.38', 'Medium', '62%', 'Opus 7/7. Sonnet 5/7. Haiku 1/7. Sharpest model-tier step function.', 'FFF3CD', '7D5A00', False),
    ('Chn — Chained', '8 / 21', '0.38', 'Medium', '62%', 'First fail: Opus T6, Sonnet T5, Haiku T4. Tracks model size directly.', 'FFF3CD', '7D5A00', False),
    ('Neg — Negation', '4 / 21', '0.19', 'Easy', '81%', 'Highest-weight Easy category. Word constraints tractable across all tiers.', 'D4EDDA', '155724', False),
    ('Per — Persona', '2 / 21', '0.10', 'Easy', '90%', 'Low failure rate across all models and tiers.', 'D4EDDA', '155724', False),
    ('Fmt / Exa / Pre', '1 / 21', '0.05', 'Easy', '95%', 'Near-zero failure rate. Mechanical rule-following well within model capacity.', 'D4EDDA', '155724', False),
    ('Cnt / Sty / Pri', '0 / 21', '0.00', 'Excluded', '100%', 'Perfect scores. No discriminative value for cross-category testing.', 'F0F0F0', '888888', False),
]
for i, row in enumerate(cat_rows):
    c1,c2,c3,c4,c5,c6,bg,txt,bold = row
    is_hdr = (i == 0)
    tc_color = 'FFFFFF' if is_hdr else txt
    styled_cell(tc.cell(i,0), c1, bg=bg, bold=True, color=tc_color)
    styled_cell(tc.cell(i,1), c2, bg=bg, bold=is_hdr, color=tc_color, align='center')
    styled_cell(tc.cell(i,2), c3, bg=bg, bold=True if (not is_hdr and i==1) else is_hdr, color=tc_color, align='center')
    styled_cell(tc.cell(i,3), c4, bg=bg, bold=True if not is_hdr else is_hdr, color=tc_color, align='center')
    styled_cell(tc.cell(i,4), c5, bg=bg, bold=is_hdr, color=tc_color, align='center')
    styled_cell(tc.cell(i,5), c6, bg=bg, bold=is_hdr, color=tc_color)

spacer(14)

# ── Section 3: Methodology ────────────────────────────────────────────
heading1('3. How Round 1 Results Determine Easy / Medium / Hard')

body('The difficulty bands are not assigned by judgment — they emerge directly from the observed failure distribution in Round 1. The five steps below show the derivation.')
spacer(6)

steps = [
    ('Step 1 — Count failures per category across all models and tiers',
     'Each category was tested at 7 tiers across 3 models = 21 possible passes per category. The raw failure count becomes the numerator of the weight formula. Categories with more platform-wide failures are harder by definition.'),
    ('Step 2 — Compute category weight (normalised failure rate)',
     'Weight = failures / 21. This converts raw counts to a 0-1 scale. A weight of 0.71 means 71% of all model-tier combinations failed that category. The scale is platform-specific — it reflects difficulty on the Anthropic API with this dataset, not a universal measure.'),
    ('Step 3 — Identify natural breaks in the weight distribution',
     'The 11 categories produce a clustered distribution. No category sits between 0.19 and 0.38, and none between 0.38 and 0.71. The thresholds at 0.20 and 0.50 fall at natural gaps in the data — the breaks are data-driven, not arbitrary cutoffs.'),
    ('Step 4 — Assign bands',
     'Hard: weight >= 0.50 (Scope only). Medium: 0.20 <= weight < 0.50 (Situation, Chained). Easy: 0 < weight < 0.20 (Negation, Persona, Format, Example, Precision). Excluded: weight = 0 (Content, Style, Priority).'),
    ('Step 5 — Apply to cross-category design',
     'In a cross-category test, bands predict relative difficulty and expected failure tier. A system prompt combining Hard + Medium categories degrades faster than one combining Easy + Medium. The weighted score formula carries this forward: failures on Hard categories penalise the overall score more than failures on Easy ones — which is why raw scores consistently overstate model performance relative to weighted scores.'),
]
for title, detail in steps:
    body(title, bold=True, size=9, space_after=2)
    body(detail, size=8.5, space_after=8)

spacer(6)

# ── Section 4: Per-model baseline ────────────────────────────────────
heading1('4. Round 1 Per-Model Performance (Baseline)')

body('Scores carried forward as the baseline for cross-category testing. Weighted scores apply platform-derived weights — failures on Hard categories count more than failures on Easy ones.')
spacer(6)

tm = doc.add_table(rows=4, cols=5)
tm.alignment = WD_TABLE_ALIGNMENT.CENTER
set_col_widths(tm, [4.5, 2.4, 2.4, 2.8, 5.9])
model_rows = [
    ('Model', 'Raw Passes', 'Raw Score', 'Weighted Score', 'Notes', '1F3864', 'FFFFFF', True),
    ('claude-opus-4-8', '71 / 77', '92%', '78%', '14-point gap: failures concentrated in Scope (Hard).', 'F2F7FC', '000000', False),
    ('claude-sonnet-4-6', '65 / 77', '84%', '52%', '32-point gap: Scope failures dominant; Situation also degraded.', 'FFFFFF', '000000', False),
    ('claude-haiku-4-5', '55 / 77', '71%', '36%', '35-point gap: failed Hard and Medium categories heavily from early tiers.', 'F2F7FC', '000000', False),
]
for i, row in enumerate(model_rows):
    c1,c2,c3,c4,c5,bg,txt,bold = row
    is_hdr = (i == 0)
    styled_cell(tm.cell(i,0), c1, bg=bg, bold=True, color='FFFFFF' if is_hdr else txt)
    styled_cell(tm.cell(i,1), c2, bg=bg, bold=is_hdr, color='FFFFFF' if is_hdr else txt, align='center')
    styled_cell(tm.cell(i,2), c3, bg=bg, bold=is_hdr, color='FFFFFF' if is_hdr else txt, align='center')
    w_color = '8B0000' if (not is_hdr and i > 1) else ('FFFFFF' if is_hdr else '000000')
    styled_cell(tm.cell(i,3), c4, bg=bg, bold=True, color=w_color, align='center')
    styled_cell(tm.cell(i,4), c5, bg=bg, bold=is_hdr, color='FFFFFF' if is_hdr else txt)

spacer(8)
body('The gap between raw and weighted score reveals how much failures concentrated in high-difficulty categories. A 14-point gap (Opus) indicates failures spread across categories. A 32-35 point gap (Sonnet, Haiku) indicates failures concentrated in Hard and Medium categories — exactly the pattern cross-category testing is designed to stress.', size=8.5)

doc.save('Round1_Adapted_Report.docx')
print('Done — Round1_Adapted_Report.docx written')
